import os
import sys
from itertools import combinations
from random import sample
import numpy as np
from docx import Document
from docx.shared import Inches
import docx.enum.text
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from PyQt5 import QtCore, QtGui, QtWidgets


class UIMainWindow(object):
    def __init__(self):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(302, 217)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.gridLayout = QtWidgets.QGridLayout(self.centralwidget)
        self.gridLayout.setObjectName("gridLayout")
        self.selectKeyLabel = QtWidgets.QLabel(self.centralwidget)
        self.selectKeyLabel.setAlignment(QtCore.Qt.AlignCenter)
        self.selectKeyLabel.setObjectName("selectKeyLabel")
        self.gridLayout.addWidget(self.selectKeyLabel, 0, 0, 1, 4)
        self.selectFileButton = QtWidgets.QPushButton(self.centralwidget)
        self.selectFileButton.setObjectName("selectFileButton")
        self.gridLayout.addWidget(self.selectFileButton, 1, 0, 1, 1)
        self.inputPathLabel = QtWidgets.QLabel(self.centralwidget)
        self.inputPathLabel.setObjectName("inputPathLabel")
        self.gridLayout.addWidget(self.inputPathLabel, 1, 1, 1, 3)
        self.possibleCardLabel = QtWidgets.QLabel(self.centralwidget)
        self.possibleCardLabel.setObjectName("possibleCardLabel")
        self.gridLayout.addWidget(self.possibleCardLabel, 2, 0, 1, 4)
        self.numCardPromptLabel = QtWidgets.QLabel(self.centralwidget)
        self.numCardPromptLabel.setObjectName("numCardPromptLabel")
        self.gridLayout.addWidget(self.numCardPromptLabel, 3, 0, 1, 2)
        self.numCardSpinBox = QtWidgets.QSpinBox(self.centralwidget)
        self.numCardSpinBox.setObjectName("numCardSpinBox")
        self.gridLayout.addWidget(self.numCardSpinBox, 3, 2, 1, 1)
        self.inputFolderPromptLabel = QtWidgets.QLabel(self.centralwidget)
        self.inputFolderPromptLabel.setAlignment(QtCore.Qt.AlignCenter)
        self.inputFolderPromptLabel.setObjectName("inputFolderPromptLabel")
        self.gridLayout.addWidget(self.inputFolderPromptLabel, 4, 0, 1, 4)
        self.outputFolderButton = QtWidgets.QPushButton(self.centralwidget)
        self.outputFolderButton.setObjectName("inputFolderButton")
        self.gridLayout.addWidget(self.outputFolderButton, 5, 0, 1, 1)
        self.outputPathLabel = QtWidgets.QLabel(self.centralwidget)
        self.outputPathLabel.setObjectName("outputPathLabel")
        self.gridLayout.addWidget(self.outputPathLabel, 5, 1, 1, 3)
        self.generateButton = QtWidgets.QPushButton(self.centralwidget)
        self.generateButton.setObjectName("generateButton")
        self.gridLayout.addWidget(self.generateButton, 6, 3, 1, 1)
        MainWindow.setCentralWidget(self.centralwidget)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)
        self.retranslate_ui(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

        self.selectFileButton.clicked.connect(self.set_file)
        self.outputFolderButton.clicked.connect(self.set_output)
        self.generateButton.clicked.connect(self.generate)

    def retranslate_ui(self, main_window):
        _translate = QtCore.QCoreApplication.translate
        main_window.setWindowTitle(_translate("MainWindow", "Bingo Generator"))
        self.selectKeyLabel.setText(_translate("MainWindow", "Select the answer key for the bingo cards you would "
                                                             "like to create:"))
        self.selectFileButton.setText(_translate("MainWindow", "Select File"))
        self.inputPathLabel.setText(_translate("MainWindow", "No file selected"))
        self.possibleCardLabel.setText(_translate("MainWindow", "This key can generate [ 0 ] different bingo cards."))
        self.numCardPromptLabel.setText(_translate("MainWindow", "How many would you like to produce?"))
        self.inputFolderPromptLabel.setText(_translate("MainWindow", "Select the directory you would like to save "
                                                                     "your cards in:"))
        self.outputFolderButton.setText(_translate("MainWindow", "Select Folder"))
        self.outputPathLabel.setText(_translate("MainWindow", "No folder selected"))
        self.generateButton.setText(_translate("MainWindow", "Generate!"))

    def set_file(self):
        input_path, _ = QtWidgets.QFileDialog.getOpenFileName(None, "Select Key", "", "Word Files (*.docx)")
        if input_path:
            self.inputPathLabel.setText(input_path)

        input_ = Document(input_path)
        table = input_.tables[0]
        problems = []
        for row in range(0, len(table.rows)):
            problems.append(table.cell(row, 0).text)

        comb = len(list(combinations(problems, 24)))
        if comb <= 1000:
            self.possibleCardLabel.setText("This key can generate [ " + str(comb) + " ] different bingo cards.")
        else:
            self.possibleCardLabel.setText("This key can generate [ more than 1000 ] different bingo cards!")

    def set_output(self):
        output_path = QtWidgets.QFileDialog.getExistingDirectory(None, "Select Directory")
        if output_path:
            self.outputPathLabel.setText(output_path)

    def generate(self):
        self.generateButton.setText("Loading...")
        input_ = Document(self.inputPathLabel.text())
        table = input_.tables[0]
        problem_set = []
        for row in range(0, len(table.rows)):
            problem_set.append(row)
        comb_full = list(combinations(problem_set, 24))
        comb = sample(comb_full, self.numCardSpinBox.value())
        comb = np.array(comb)

        document = Document('template.docx')

        for i in range(0, self.numCardSpinBox.value()):
            table = document.add_table(rows=5, cols=5)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'TableGrid'
            table.autofit = False
            for ind in range(0, 5):
                table.rows[ind].height = Inches(1.2)
                table.columns[ind].width = Inches(1.2)
            center_cell = table.cell(2, 2)
            p = center_cell.add_paragraph("FREE SPACE")
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            center_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            np.random.shuffle(comb[i])

            count = 0
            for row in range(0, 5):
                for col in range(0, 5):
                    if row != 2 or col != 2:
                        cell = table.cell(row, col)
                        copy_cell(cell, input_.tables[0].cell(comb[i][count], 0).paragraphs[0])
                        cell.paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        count += 1
            document.add_page_break()

        if os.path.exists(self.outputPathLabel.text()):
            document.save(self.outputPathLabel.text() + "/bingo_cards.docx")
        else:
            self.generateButton.setText("Path does not exist!")
        self.generateButton.setText("Done!")


def copy_cell(output_cell, input_paragraph):
    output_para = output_cell.paragraphs[0]
    for run in input_paragraph.runs:
        output_run = output_para.add_run(run.text)
        output_run.font.superscript = run.font.superscript
        output_run.font.subscript = run.font.subscript
        output_run.bold = run.bold
        output_run.italic = run.italic
        output_run.underline = run.underline
        output_run.font.color.rgb = run.font.color.rgb
        output_run.style.name = run.style.name


if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)
    MainWindow = QtWidgets.QMainWindow()
    ui = UIMainWindow()
    MainWindow.show()
    sys.exit(app.exec_())
